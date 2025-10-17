#!/usr/bin/env python3
import argparse
import sys
import re
from pathlib import Path
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from googletrans import Translator

def set_table_borders_black(table):
    """Apply black borders to the Word table."""
    tbl_elem = table._tbl
    tbl_pr = tbl_elem.xpath("./w:tblPr")[0]
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_elem = OxmlElement(f"w:{border_name}")
        border_elem.set(qn('w:val'), 'single')
        border_elem.set(qn('w:sz'), '4')  # Thin line
        border_elem.set(qn('w:space'), '0')
        border_elem.set(qn('w:color'), '000000')  # Black color
        tbl_borders.append(border_elem)
    tbl_pr.append(tbl_borders)

def parse_srt(srt_content):
    """Extract timecodes and text from SRT content."""
    pattern = re.compile(
        r'(\d+)\s+(\d{2}:\d{2}:\d{2},\d{3}) --> (\d{2}:\d{2}:\d{2},\d{3})\s+(.*?)\s*(?=\n\d+\n|\Z)',
        re.DOTALL
    )
    return pattern.findall(srt_content)

def create_docx(blocks, output_path, langs):
    """Generate the DOCX file with a table."""
    doc = Document()
    doc.add_heading('Dialogue List', level=1)

    # Table columns: Timecode + len(langs) + English
    num_cols = 2 + len(langs)
    table = doc.add_table(rows=1, cols=num_cols)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Timecode'
    for idx, lang in enumerate(langs):
        hdr_cells[1 + idx].text = lang.capitalize()
    hdr_cells[-1].text = 'English'

    translator = Translator() if langs else None

    for _, start, end, text in blocks:
        start_time = start.split(',')[0]
        end_time = end.split(',')[0]
        timecode = f"{start_time} --> {end_time}"
        subtitle_text = ' '.join(text.strip().splitlines())

        row_cells = table.add_row().cells
        row_cells[0].text = timecode

        for idx, lang_code in enumerate(langs):
            if translator:
                try:
                    translated = translator.translate(subtitle_text, dest=lang_code).text
                except Exception:
                    translated = ''
            else:
                translated 
            row_cells[1 + idx].text = translated
        row_cells[-1].text = subtitle_text

    set_table_borders_black(table)
    doc.save(output_path)

def main():
    parser = argparse.ArgumentParser(description="Convert SRT to DOCX with optional translations")
    parser.add_argument('input', help='Path to input .srt file')
    parser.add_argument('output', nargs='?', help='Path to output .docx file (optional)')
    parser.add_argument('--langs', nargs='*', default=[], help='Language codes to translate to (e.g. de fr es)')
    
    args = parser.parse_args()
    
    input_path = Path(args.input)
    if not args.output:
        output_path = input_path.with_suffix('.docx')
    else:
        output_path = Path(args.output)

    # Read SRT
    with input_path.open('r', encoding='utf-8') as f:
        srt_content = f.read()
    
    subtitles = parse_srt(srt_content)
    create_docx(subtitles, output_path, args.langs)

    print(f"✅ DOCX file created: {output_path}")

if __name__ == "__main__":
    main()
